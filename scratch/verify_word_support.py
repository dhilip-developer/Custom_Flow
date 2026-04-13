import io
import asyncio
from services.pipeline.splitter import DocumentSplitter

# We need python-docx to generate a test file
try:
    from docx import Document
except ImportError:
    import subprocess
    subprocess.run(["pip", "install", "python-docx"], check=True)
    from docx import Document

def generate_test_docx():
    doc = Document()
    doc.add_heading('Customs Intelligence Test', 0)
    doc.add_paragraph('This is a test invoice document for Agent 1.')
    doc.add_paragraph('Invoice No: INV-WORD-001')
    doc.add_paragraph('Total: 5000 USD')
    
    # Save to buffer
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def verify_word_support():
    print("Generating test .docx file...")
    docx_bytes = generate_test_docx()
    print(f"Generated {len(docx_bytes)} bytes of docx.")
    
    print("\nExecuting Pipeline Splitter...")
    try:
        pages = DocumentSplitter.split(docx_bytes)
        
        print(f"\n[SUCCESS] Splitter returned {len(pages)} pages.")
        for p in pages:
            print(f"Page {p['page_number']}:")
            print(f"  - Text Length: {len(p['text'])} chars")
            print(f"  - Image (JPG) size: {len(p['image_bytes'])} bytes")
            print(f"  - PDF Buffer size: {len(p['buffer'])} bytes")
            
            if "INV-WORD-001" in p['text']:
                print("  - [VERIFIED] Found specific invoice text in extracted data.")
                
    except Exception as e:
        print(f"\n[FAILED] Splitter error: {e}")

if __name__ == "__main__":
    verify_word_support()
